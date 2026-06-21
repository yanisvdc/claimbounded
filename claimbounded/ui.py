"""Interactive Gradio UI for claimbounded.

Launch with:
    claimbounded ui
or from Python:
    from claimbounded.ui import launch
    launch()
"""

from __future__ import annotations

import base64
import tempfile
from datetime import date as _date
from typing import Any

# ---------------------------------------------------------------------------
# Controlled vocabulary
# ---------------------------------------------------------------------------

_UNCLEAR = "unclear"

SUBMISSION_PATHWAYS = ["510k", "de_novo", _UNCLEAR]
CLINICAL_DOMAINS = [
    "cardiology", "dental", "general_hospital", "laboratory_ivd",
    "neurology", "oncology", "ophthalmology", "other",
    "pathology", "radiology", "surgery", _UNCLEAR,
]
DEVICE_FUNCTIONS = [
    "triage_notification", "diagnostic_classification", "segmentation_quantification",
    "risk_prediction_prognosis", "physiologic_monitoring", "acquisition_guidance",
    "image_reconstruction_enhancement", "therapy_planning_or_dosing",
    "pacs_viewer_workflow", "workflow_planning_guidance",
    "laboratory_ivd_classification", "other", _UNCLEAR,
]
INPUT_DATA_TYPES = [
    "radiology_image", "pathology_image", "physiologic_signal",
    "clinical_record_data", "laboratory_measurement", "device_sensor_data",
    "genomic_or_molecular_data", "mixed", "other", _UNCLEAR,
]
ENDPOINT_TYPES = [
    "diagnostic_accuracy", "triage_sensitivity_specificity",
    "quantitative_measurement_agreement", "segmentation_geometric_accuracy",
    "image_quality_or_reconstruction_fidelity", "risk_prediction_or_prognosis",
    "physiologic_event_detection", "workflow_or_time_to_notification",
    "technical_performance_only", "substantial_equivalence_only", _UNCLEAR,
]
GROUND_TRUTH_MODALITIES = [
    "expert_reader_panel", "expert_annotation", "clinical_diagnosis",
    "laboratory_reference_method", "pathology_or_histology",
    "longitudinal_clinical_outcome", "physiologic_reference_standard",
    "phantom_or_bench_reference", "predicate_device_comparison",
    "not_reported", _UNCLEAR,
]
EVIDENCE_STREAMS = [
    "workflow_logs", "human_corrections_or_edits", "clinician_acceptance_or_override",
    "output_logs_only", "technical_logs_only", "downstream_clinical_outcome",
    "none_described", _UNCLEAR,
]
YES_NO_UNCLEAR = ["yes", "no", _UNCLEAR]
RETRIEVAL_MODES = ["hybrid", "like_for_like", "adjacent", "claim_gap"]

_CLAIM_HIERARCHY = [
    "no_performance_claim_auditable", "utilization_only",
    "technical_pipeline_stability", "workflow_performance",
    "human_machine_concordance", "output_quality_or_measurement_agreement",
    "clinical_accuracy_or_calibration",
]
_CLAIM_RANK = {c: i for i, c in enumerate(_CLAIM_HIERARCHY)}
_CLAIM_LABELS = {
    "no_performance_claim_auditable": "No performance claim auditable",
    "utilization_only": "Utilization only",
    "technical_pipeline_stability": "Technical pipeline stability",
    "workflow_performance": "Workflow performance",
    "human_machine_concordance": "Human–machine concordance",
    "output_quality_or_measurement_agreement": "Output quality / measurement agreement",
    "clinical_accuracy_or_calibration": "Clinical accuracy or calibration",
}
_ENDPOINT_TO_CLAIM = {
    "diagnostic_accuracy": "clinical_accuracy_or_calibration",
    "triage_sensitivity_specificity": "clinical_accuracy_or_calibration",
    "risk_prediction_or_prognosis": "clinical_accuracy_or_calibration",
    "physiologic_event_detection": "clinical_accuracy_or_calibration",
    "quantitative_measurement_agreement": "output_quality_or_measurement_agreement",
    "segmentation_geometric_accuracy": "output_quality_or_measurement_agreement",
    "image_quality_or_reconstruction_fidelity": "output_quality_or_measurement_agreement",
    "technical_performance_only": "technical_pipeline_stability",
    "substantial_equivalence_only": "technical_pipeline_stability",
    "workflow_or_time_to_notification": "workflow_performance",
}

# ---------------------------------------------------------------------------
# Example + clear defaults
# ---------------------------------------------------------------------------

_EXAMPLE_FIELD_ORDER = [
    "device_name", "applicant", "submission_number", "submission_pathway",
    "clinical_domain", "device_function", "input_data_type",
    "authorization_endpoint_type", "authorization_ground_truth_modality",
    "routine_postmarket_evidence_stream", "endpoint_linked_to_ai_output",
    "endpoint_routinely_recorded", "human_correction_available",
    "human_overread_or_confirmation_required",
]
_EXAMPLE = dict(
    device_name="Acme LVO Triage",
    applicant="Acme Medical Inc.",
    submission_number="",
    submission_pathway="510k",
    clinical_domain="neurology",
    device_function="triage_notification",
    input_data_type="radiology_image",
    authorization_endpoint_type="diagnostic_accuracy",
    authorization_ground_truth_modality="expert_reader_panel",
    routine_postmarket_evidence_stream="workflow_logs",
    endpoint_linked_to_ai_output="no",
    endpoint_routinely_recorded="no",
    human_correction_available="no",
    human_overread_or_confirmation_required="no",
)
_CLEAR = dict(
    device_name="", applicant="", submission_number="",
    submission_pathway=_UNCLEAR, clinical_domain=_UNCLEAR,
    device_function=_UNCLEAR, input_data_type=_UNCLEAR,
    authorization_endpoint_type=_UNCLEAR,
    authorization_ground_truth_modality=_UNCLEAR,
    routine_postmarket_evidence_stream=_UNCLEAR,
    endpoint_linked_to_ai_output=_UNCLEAR,
    endpoint_routinely_recorded=_UNCLEAR,
    human_correction_available=_UNCLEAR,
    human_overread_or_confirmation_required=_UNCLEAR,
)

# ---------------------------------------------------------------------------
# HTML report CSS  (sky-blue/teal palette, scoped to .cbr-wrap)
# ---------------------------------------------------------------------------

_CSS_SCOPED = """
.cbr-wrap{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;
  color:#0f172a;line-height:1.65;background:#f0f9ff;}
.cbr-wrap *{box-sizing:border-box;}
.cbr-container{max-width:980px;margin:0 auto;padding:48px 32px;}
.cbr-header{background:linear-gradient(135deg,#0369a1 0%,#0891b2 100%);color:white;
  padding:40px 48px;border-radius:16px;margin-bottom:28px;}
.cbr-brand{font-size:11px;font-weight:700;letter-spacing:2px;text-transform:uppercase;
  color:#bae6fd;margin-bottom:12px;}
.cbr-header h1{font-size:28px;font-weight:800;margin-bottom:8px;color:white;}
.cbr-meta{font-size:13px;opacity:.85;color:white;}
.cbr-stats{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-bottom:24px;}
.cbr-stat{background:white;border-radius:12px;padding:22px;text-align:center;
  box-shadow:0 1px 4px rgba(0,0,0,.07);border-top:3px solid #0891b2;}
.cbr-val{font-size:26px;font-weight:800;color:#0369a1;}
.cbr-lbl{font-size:12px;color:#64748b;margin-top:4px;}
.cbr-callout{border-radius:12px;padding:28px 32px;margin-bottom:24px;color:white;}
.cbr-green{background:linear-gradient(135deg,#16a34a,#059669);}
.cbr-amber{background:linear-gradient(135deg,#b45309,#92400e);}
.cbr-red{background:linear-gradient(135deg,#dc2626,#b91c1c);}
.cbr-clabel{font-size:11px;font-weight:700;letter-spacing:1.5px;text-transform:uppercase;
  opacity:.85;margin-bottom:6px;}
.cbr-cvalue{font-size:26px;font-weight:800;margin-bottom:6px;}
.cbr-csub{font-size:14px;opacity:.95;}
.cbr-section{background:white;border-radius:12px;padding:32px;margin-bottom:20px;
  box-shadow:0 1px 4px rgba(0,0,0,.07);}
.cbr-section h2{font-size:17px;font-weight:700;color:#0369a1;margin-bottom:18px;
  padding-bottom:12px;border-bottom:2px solid #bae6fd;}
.cbr-ladder{display:flex;flex-direction:column;gap:6px;}
.cbr-rung{display:flex;align-items:center;padding:11px 16px;border-radius:8px;
  font-size:14px;gap:10px;}
.cbr-ceiling{background:#16a34a;color:white;font-weight:700;}
.cbr-supported{background:#dcfce7;color:#166534;border:1px solid #86efac;}
.cbr-auth-gap{background:#fef9c3;color:#854d0e;border:2px dashed #fbbf24;}
.cbr-unsupported{background:#f8fafc;color:#94a3b8;border:1px solid #e2e8f0;}
.cbr-icon{font-size:15px;width:20px;text-align:center;flex-shrink:0;}
.cbr-rl{flex:1;}
.cbr-rb{font-size:11px;padding:2px 10px;border-radius:10px;
  background:rgba(255,255,255,.3);font-weight:600;}
.cbr-rn{font-size:11px;color:#92400e;font-weight:600;}
.cbr-tw{overflow-x:auto;}
.cbr-section table{width:100%;border-collapse:collapse;font-size:13px;}
.cbr-section th{background:#e0f2fe;color:#0369a1;font-weight:600;font-size:11px;
  text-transform:uppercase;letter-spacing:.4px;padding:10px 12px;text-align:left;
  border-bottom:2px solid #bae6fd;}
.cbr-section td{padding:10px 12px;border-bottom:1px solid #f0f9ff;vertical-align:top;}
.cbr-section tr:last-child td{border-bottom:none;}
.cbr-yes{color:#16a34a;font-weight:700;}
.cbr-partial{color:#d97706;font-weight:600;}
.cbr-no{color:#94a3b8;}
.cbr-2col{display:grid;grid-template-columns:1fr 1fr;gap:20px;margin-bottom:20px;}
.cbr-info{background:#f0f9ff;border-left:4px solid #0891b2;padding:14px 18px;
  border-radius:0 8px 8px 0;margin-bottom:12px;}
.cbr-info.am{border-left-color:#d97706;background:#fffbeb;}
.cbr-ilabel{font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:1px;
  color:#64748b;margin-bottom:3px;}
.cbr-ivalue{font-size:14px;font-weight:600;color:#0f172a;}
.cbr-isub{font-size:13px;color:#475569;margin-top:3px;}
.cbr-list{list-style:none;padding:0;}
.cbr-list li{padding:8px 0 8px 22px;border-bottom:1px solid #f0f9ff;
  position:relative;font-size:13px;}
.cbr-list li:last-child{border-bottom:none;}
.cbr-list li::before{content:'→';position:absolute;left:0;color:#0891b2;font-weight:700;}
.cbr-olist{list-style:none;padding:0;counter-reset:item;}
.cbr-olist li{counter-increment:item;padding:8px 0 8px 30px;
  border-bottom:1px solid #f0f9ff;position:relative;font-size:13px;}
.cbr-olist li:last-child{border-bottom:none;}
.cbr-olist li::before{content:counter(item);position:absolute;left:0;color:white;
  background:#0891b2;font-size:10px;font-weight:700;width:18px;height:18px;
  border-radius:50%;display:flex;align-items:center;justify-content:center;top:10px;}
.cbr-mnote{font-size:12px;color:#64748b;background:#f0f9ff;border-radius:8px;
  padding:12px 16px;margin-bottom:16px;border:1px solid #bae6fd;}
.cbr-footer{background:#e0f2fe;border-radius:12px;padding:22px 28px;margin-top:28px;
  border-left:4px solid #0891b2;}
.cbr-footer p{font-size:12px;color:#475569;line-height:1.65;}
.cbr-wrap code{background:#e0f2fe;padding:1px 5px;border-radius:4px;
  font-family:monospace;font-size:12px;color:#0369a1;}
@media(max-width:640px){.cbr-stats{grid-template-columns:1fr;}
  .cbr-2col{grid-template-columns:1fr;}.cbr-container{padding:24px 16px;}}
"""
_CSS_STANDALONE = f"* {{box-sizing:border-box;margin:0;padding:0;}}\nbody {{background:#f0f9ff;}}\n{_CSS_SCOPED}"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _humanize_match(match_str: str) -> str:
    if not match_str or match_str == "text similarity only":
        return "Text similarity only"
    labels = {
        "product_code": "product code", "device_function": "device function",
        "authorization_endpoint_type": "auth. endpoint", "authorization_endpoint": "auth. endpoint",
        "authorization_ground_truth_modality": "ground truth", "ground_truth": "ground truth",
        "routine_postmarket_evidence_stream": "evidence stream", "evidence_stream": "evidence stream",
        "strongest_auditable_postmarket_claim": "claim ceiling", "claim_ceiling": "claim ceiling",
        "postmarket_audit_burden": "audit burden",
    }
    out = []
    for p in match_str.replace("shares ", "").split("; "):
        if "=" in p:
            k, v = p.split("=", 1)
            out.append(f"Same {labels.get(k.strip(), k.strip().replace('_',' '))} ({v.strip().replace('_',' ')})")
        else:
            out.append(p.strip())
    return "; ".join(out) if out else "Text similarity"


def _v(d: dict, key: str) -> str:
    """Return a field value with underscores replaced, or empty string."""
    v = d.get(key, "")
    return v.replace("_", " ") if v and v != "unclear" else ""


def _format_device_panel(profile: Any, extra: dict | None = None) -> str:
    """Narrative device profile for click-to-expand rows."""
    if profile is None:
        return "_No corpus record found for this submission number._"
    d = profile.to_dict()
    sub = d.get("submission_number", "")
    pathway = d.get("submission_pathway", "").replace("510k", "510(k)").replace("de_novo", "De Novo")
    name = d.get("device_name", "Device")
    applicant = d.get("applicant", "")
    year = d.get("year", "")

    lines = [f"#### {name}"]
    meta_parts = [p for p in [applicant, f"{pathway} {sub}".strip(), year] if p and p != "unclear"]
    lines += [" · ".join(meta_parts), ""]

    # Quick-scan row
    tags = []
    for key in ("clinical_domain", "device_function", "input_data_type"):
        v = _v(d, key)
        if v:
            tags.append(f"`{v}`")
    if tags:
        lines += [" · ".join(tags), ""]

    # Authorization & evidence
    ep = _v(d, "authorization_endpoint_type")
    gt = _v(d, "authorization_ground_truth_modality")
    stream = _v(d, "routine_postmarket_evidence_stream")
    ceiling = _v(d, "strongest_auditable_postmarket_claim")
    burden = _v(d, "postmarket_audit_burden")

    if ep or gt:
        lines += ["---", "**What it was authorized on**", ""]
        if ep:
            lines.append(f"Authorization endpoint: **{ep}**")
        if gt:
            lines.append(f"Ground truth: {gt}")
        lines.append("")

    if stream or ceiling:
        lines += ["---", "**What routine deployment generates**", ""]
        if stream:
            lines.append(f"Routine evidence stream: {stream}")
        if ceiling:
            lines.append(f"Postmarket claim ceiling: **{ceiling}**")
        if burden:
            lines.append(f"Audit burden: {burden}")
        lines.append("")

    # Full intended use
    iu = d.get("intended_use_summary", "")
    if iu and iu not in ("", "unclear"):
        lines += ["---", "**Full intended use**", "", iu, ""]

    # Authorization performance claim
    apc = d.get("authorization_performance_claim", "")
    if apc and apc not in ("", "unclear"):
        lines += ["---", "**Authorization performance claim**", "", apc, ""]

    # Supporting quote
    sq = d.get("supporting_quote_authorization", "")
    if sq and sq not in ("", "unclear"):
        lines += [
            "---",
            f"> {sq}",
            f">",
            f"> *— Public FDA {pathway} Summary {sub}*",
            "",
        ]

    if sub:
        lines += ["---", f"*Source: accessdata.fda.gov · {pathway} {sub}*"]

    return "\n".join(lines)


def _profile_card_html(d: dict, idx: int) -> str:
    """Render one device as an HTML card for the profile report."""
    sub = d.get("submission_number", "")
    pathway = d.get("submission_pathway", "").replace("510k", "510(k)").replace("de_novo", "De Novo")
    name = d.get("device_name", "")
    applicant = d.get("applicant", "")
    year = d.get("year", "")
    iu = d.get("intended_use_summary", "")
    apc = d.get("authorization_performance_claim", "")
    sq = d.get("supporting_quote_authorization", "")
    ceiling = d.get("strongest_auditable_postmarket_claim", "").replace("_", " ")
    burden = d.get("postmarket_audit_burden", "").replace("_", " ")

    def field(label, key):
        v = d.get(key, "")
        if v and v != "unclear":
            return (f'<div class="cbr-info" style="margin-bottom:8px">'
                    f'<div class="cbr-ilabel">{label}</div>'
                    f'<div class="cbr-ivalue" style="font-size:13px">{v.replace("_"," ")}</div></div>')
        return ""

    fields_html = "".join([
        field("Clinical domain", "clinical_domain"),
        field("Device function", "device_function"),
        field("Input data type", "input_data_type"),
        field("Authorization endpoint type", "authorization_endpoint_type"),
        field("Ground truth modality", "authorization_ground_truth_modality"),
        field("Routine evidence stream", "routine_postmarket_evidence_stream"),
    ])

    ceiling_color = "#16a34a" if "accuracy" in ceiling or "calibration" in ceiling else (
        "#d97706" if "concordance" in ceiling or "quality" in ceiling else "#0891b2"
    )

    iu_section = (f'<div style="margin:16px 0">'
                  f'<div class="cbr-ilabel">Full intended use</div>'
                  f'<p style="margin-top:8px;font-size:13px;line-height:1.7;color:#0f172a">{iu}</p>'
                  f'</div>') if iu and iu != "unclear" else ""

    apc_section = (f'<div style="margin:16px 0">'
                   f'<div class="cbr-ilabel">Authorization performance claim</div>'
                   f'<p style="margin-top:8px;font-size:13px;line-height:1.7;color:#0f172a">{apc}</p>'
                   f'</div>') if apc and apc != "unclear" else ""

    quote_section = (f'<blockquote style="margin:16px 0;padding:12px 16px;border-left:4px solid #bae6fd;'
                     f'background:#f0f9ff;border-radius:0 8px 8px 0;font-style:italic;font-size:13px;color:#475569">'
                     f'"{sq}"<br><small style="color:#94a3b8">— Public FDA {pathway} Summary {sub}</small>'
                     f'</blockquote>') if sq and sq != "unclear" else ""

    ceiling_badge = (f'<div style="display:inline-block;padding:4px 12px;border-radius:20px;'
                     f'background:{ceiling_color};color:white;font-size:12px;font-weight:700;margin-top:8px">'
                     f'{ceiling}</div>') if ceiling and ceiling != "unclear" else ""

    return f"""
<div class="cbr-section" style="margin-bottom:24px">
  <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:8px">
    <div>
      <h2 style="border-bottom:none;margin-bottom:4px;padding-bottom:0">{idx}. {name}</h2>
      <div style="color:#64748b;font-size:13px">{applicant} · {pathway} <code>{sub}</code> · {year}</div>
    </div>
    {ceiling_badge}
  </div>
  <hr style="border:none;border-top:2px solid #bae6fd;margin:16px 0">
  <div class="cbr-2col" style="margin-bottom:0">{fields_html}</div>
  {f'<div class="cbr-info am"><div class="cbr-ilabel">Postmarket claim ceiling</div><div class="cbr-ivalue">{ceiling}</div><div class="cbr-isub">Audit burden: {burden}</div></div>' if ceiling and ceiling != "unclear" else ""}
  {iu_section}
  {apc_section}
  {quote_section}
  <div style="font-size:11px;color:#94a3b8;margin-top:8px">Source: accessdata.fda.gov · {pathway} {sub}</div>
</div>"""


def _build_profile_html(profiles: list, title: str = "Claim-Bounded Monitoring of AI-Enabled Medical Devices") -> str:
    """Build a beautiful HTML string for search results or lookup (no file I/O)."""
    today = str(_date.today())
    n = len(profiles)

    cards = "".join(
        _profile_card_html(p.to_dict() if hasattr(p, "to_dict") else p, i)
        for i, p in enumerate(profiles, 1)
    )

    html_str = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>claimbounded — {title}</title>
<style>{_CSS_STANDALONE}</style>
</head>
<body>
<div class="cbr-wrap"><div class="cbr-container">

<header class="cbr-header">
  <div class="cbr-brand">claimbounded · FDA AI Device Corpus · schema v3_auditability</div>
  <h1>{title}</h1>
  <div class="cbr-meta">{n} device{"s" if n != 1 else ""} · {today} · Grounded in 1,404 public FDA authorization records</div>
</header>

<div class="cbr-section">
  <h2>How to Use This Report</h2>
  <div class="cbr-2col">
    <div>
      <div style="font-size:22px;margin-bottom:8px">📋</div>
      <div style="font-weight:700;margin-bottom:6px;color:#0369a1">Regulators</div>
      <div style="font-size:13px;color:#475569;line-height:1.6">
        These public authorization records reveal what evidence comparable devices provided at authorization,
        what reference standard they used, and what postmarket claim ceiling their deployment evidence supports.
        Cross-reference submission numbers at <strong>accessdata.fda.gov</strong> for the full authorization summary.
      </div>
    </div>
    <div>
      <div style="font-size:22px;margin-bottom:8px">🔬</div>
      <div style="font-weight:700;margin-bottom:6px;color:#0369a1">Device Manufacturers</div>
      <div style="font-size:13px;color:#475569;line-height:1.6">
        Benchmark your evidence architecture against comparable authorized devices.
        The <em>claim ceiling</em> shows the strongest postmarket monitoring claim each device can support
        from routine deployment data — a target to design toward in your own logging and export features.
      </div>
    </div>
  </div>
  <div style="margin-top:20px">
    <div style="font-size:22px;margin-bottom:8px">🏥</div>
    <div style="font-weight:700;margin-bottom:6px;color:#0369a1">Health Systems</div>
    <div style="font-size:13px;color:#475569;line-height:1.6">
      Before procuring an AI device, use these profiles to understand what comparable devices
      can and cannot substantiate about postmarket performance. Focus on the <em>claim ceiling</em>
      and <em>audit burden</em> — they tell you the strongest monitoring claim the device can support
      from routine data, and how much additional evidence work closing that gap would require.
    </div>
  </div>
</div>

{cards}

<div class="cbr-footer">
  <p><strong>Grounding note:</strong> All profiles are drawn from publicly available FDA 510(k) and De Novo
  authorization summaries for AI-enabled medical devices, coded under the claimbounded study codebook
  (schema v3_auditability). This document does not constitute a regulatory determination, does not predict
  FDA decisions, and does not assess device safety or effectiveness. It is intended solely to support
  evidence planning, procurement evaluation, and regulatory review.</p>
</div>

</div></div></body></html>"""

    return html_str


def _save_profile_html(html_str: str) -> str:
    tmp = tempfile.NamedTemporaryFile(mode="w", suffix="_claimbounded_profiles.html", delete=False, encoding="utf-8")
    tmp.write(html_str)
    tmp.close()
    return tmp.name


def _generate_profile_docx(profiles: list, title: str = "Device Profile Report") -> str | None:
    """Generate a Word document for search results or lookup. Returns temp file path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    today = str(_date.today())

    def _bg(cell, color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color.lstrip("#"))
        tcPr.append(shd)

    def _run(para, text: str, bold=False, color: str | None = None,
             size: float | None = None, italic=False):
        r = para.add_run(text)
        r.bold = bold
        r.italic = italic
        if color:
            rv, gv, bv = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
            r.font.color.rgb = RGBColor(rv, gv, bv)
        if size:
            r.font.size = Pt(size)
        return r

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)

    h0 = doc.add_heading(level=0)
    _run(h0, "claimbounded", color="#0369a1", size=20, bold=True)
    h1 = doc.add_heading(level=1)
    h1.add_run(title)
    p = doc.add_paragraph()
    _run(p, f"{len(profiles)} device(s) · {today} · Grounded in 1,404 public FDA authorization records",
         italic=True, color="#64748b", size=10)
    doc.add_paragraph()

    # Stakeholder guidance box
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _bg(cell, "#e0f2fe")
    sp = cell.paragraphs[0]
    sp.paragraph_format.space_before = Pt(6)
    _run(sp, "HOW TO USE THIS REPORT\n", bold=True, color="#0369a1", size=9)
    _run(sp, "📋 Regulators: ", bold=True, color="#0369a1", size=10)
    _run(sp, "Cross-reference submission numbers at accessdata.fda.gov. Review the claim ceiling and audit burden to assess whether marketed monitoring claims are supportable.\n", color="#475569", size=10)
    _run(sp, "🔬 Manufacturers: ", bold=True, color="#0f766e", size=10)
    _run(sp, "Benchmark your evidence architecture. The claim ceiling shows the strongest postmarket claim comparable authorized devices can support from routine data.\n", color="#475569", size=10)
    _run(sp, "🏥 Health Systems: ", bold=True, color="#6b21a8", size=10)
    _run(sp, "Focus on claim ceiling and audit burden before procurement. These determine what monitoring claims the device can substantiate without additional evidence work.", color="#475569", size=10)
    sp.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()

    for i, profile in enumerate(profiles, 1):
        d = profile.to_dict() if hasattr(profile, "to_dict") else profile
        sub = d.get("submission_number", "")
        pathway = d.get("submission_pathway", "").replace("510k", "510(k)").replace("de_novo", "De Novo")
        name = d.get("device_name", "Device")

        doc.add_heading(f"{i}. {name}", 2)
        meta_p = doc.add_paragraph()
        parts = [x for x in [d.get("applicant",""), f"{pathway} {sub}".strip(), d.get("year","")] if x and x != "unclear"]
        _run(meta_p, " · ".join(parts), color="#64748b", size=10)

        for label, key in [
            ("Clinical domain", "clinical_domain"), ("Device function", "device_function"),
            ("Authorization endpoint type", "authorization_endpoint_type"),
            ("Ground truth modality", "authorization_ground_truth_modality"),
            ("Routine evidence stream", "routine_postmarket_evidence_stream"),
            ("Postmarket claim ceiling", "strongest_auditable_postmarket_claim"),
            ("Audit burden", "postmarket_audit_burden"),
        ]:
            v = d.get(key, "")
            if v and v != "unclear":
                fp = doc.add_paragraph()
                _run(fp, f"{label}: ", bold=True, size=10)
                _run(fp, v.replace("_", " "), size=10)

        iu = d.get("intended_use_summary", "")
        if iu and iu != "unclear":
            doc.add_paragraph()
            doc.add_heading("Intended Use", 3)
            doc.add_paragraph(iu).runs[0].font.size = Pt(10)

        apc = d.get("authorization_performance_claim", "")
        if apc and apc != "unclear":
            doc.add_heading("Authorization Performance Claim", 3)
            doc.add_paragraph(apc).runs[0].font.size = Pt(10)

        sq = d.get("supporting_quote_authorization", "")
        if sq and sq != "unclear":
            doc.add_heading("Supporting Quote (FDA Authorization Summary)", 3)
            qp = doc.add_paragraph()
            _run(qp, f'"{sq}"', italic=True, color="#475569", size=10)
            qp2 = doc.add_paragraph()
            _run(qp2, f"— Public FDA {pathway} Summary {sub}", color="#94a3b8", size=9)

        if sub:
            sp2 = doc.add_paragraph()
            _run(sp2, f"Source: accessdata.fda.gov · {pathway} {sub}", italic=True, color="#94a3b8", size=9)

        if i < len(profiles):
            doc.add_paragraph()
            doc.add_paragraph("─" * 60).runs[0].font.color.rgb = RGBColor(0xba, 0xe6, 0xfd)
            doc.add_paragraph()

    doc.add_paragraph()
    dp = doc.add_paragraph()
    _run(dp, "Grounding note: ", bold=True, color="#475569")
    _run(dp, "All profiles are drawn from publicly available FDA authorization summaries under the claimbounded study codebook (schema v3_auditability). This document does not constitute a regulatory determination.", italic=True, color="#64748b", size=9)

    tmp = tempfile.NamedTemporaryFile(suffix="_claimbounded_profiles.docx", delete=False)
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


def _generate_html_report(pkg: dict[str, Any]) -> str:
    today = str(_date.today())
    cp = pkg["claim_profile"]
    device = pkg["device"]
    rem = cp["authorization_remeasurement"]
    ceiling = cp["routine_evidence_claim_ceiling"]
    ceiling_label = _CLAIM_LABELS.get(ceiling, ceiling)
    auth_endpoint_type = device.get("authorization_endpoint_type", _UNCLEAR)
    auth_claim_key = _ENDPOINT_TO_CLAIM.get(auth_endpoint_type)
    ceiling_rank = _CLAIM_RANK.get(ceiling, -1)
    gap_levels = rem.get("claim_gap_levels", 0)
    can_audit = rem.get("can_audit_authorization_endpoint_with_routine_data", "")
    device_name = device.get("device_name") or "Unnamed device"
    applicant = device.get("applicant", "")
    submission = device.get("submission_number", "")
    dash = pkg["dashboard_claim_limits"]

    ladder = []
    for claim in reversed(_CLAIM_HIERARCHY):
        rank = _CLAIM_RANK[claim]
        label = _CLAIM_LABELS[claim]
        is_ceiling = claim == ceiling
        is_gap = claim == auth_claim_key and auth_claim_key != ceiling and rank > ceiling_rank
        is_ok = rank < ceiling_rank
        if is_ceiling:
            cls, icon, badge = "cbr-rung cbr-ceiling", "▲", '<span class="cbr-rb">CLAIM CEILING</span>'
        elif is_gap:
            cls, icon, badge = "cbr-rung cbr-auth-gap", "⚠", '<span class="cbr-rn">Authorization target — gap exists</span>'
        elif is_ok:
            cls, icon, badge = "cbr-rung cbr-supported", "✓", ""
        else:
            cls, icon, badge = "cbr-rung cbr-unsupported", "–", ""
        ladder.append(f'<div class="{cls}"><span class="cbr-icon">{icon}</span><span class="cbr-rl">{label}</span>{badge}</div>')

    matrix_rows = []
    for row in pkg["claim_support_matrix"]:
        v = row.get("supported_by_routine_evidence", "")
        vc = "cbr-yes" if v == "Yes" else ("cbr-partial" if "link" in v.lower() else "cbr-no")
        matrix_rows.append(f'<tr><td>{row["claim"]}</td><td class="{vc}">{v}</td><td>{row["evidence_needed"]}</td></tr>')

    prec_rows = []
    for i, p in enumerate(pkg["precedents"], 1):
        use = str(p.get("intended_use_summary", ""))
        pway = p.get("submission_pathway", "").replace("510k", "510(k)").replace("de_novo", "De Novo")
        match = _humanize_match(p.get("match", ""))
        prec_rows.append(
            f'<tr><td><strong>{i}</strong></td>'
            f'<td><code>{p.get("submission_number","")}</code><br>'
            f'<span style="font-size:11px;color:#64748b">{pway} · {p.get("year","")}</span></td>'
            f'<td>{p.get("device_name","")}<br>'
            f'<span style="font-size:11px;color:#64748b">{p.get("applicant","")}</span></td>'
            f'<td style="font-size:12px;color:#475569">{use}</td>'
            f'<td>{p.get("authorization_endpoint_type","").replace("_"," ")}</td>'
            f'<td><strong>{p.get("strongest_auditable_postmarket_claim","").replace("_"," ")}</strong></td>'
            f'<td>{round(float(p.get("score",0)),3)}</td>'
            f'<td style="font-size:12px;color:#64748b">{match}</td></tr>'
        )

    n_support = len(cp.get("supportable_claims", []))
    n_prec = len(pkg.get("precedents", []))
    callout_cls = "cbr-callout cbr-green" if gap_levels <= 0 else ("cbr-callout cbr-amber" if gap_levels == 1 else "cbr-callout cbr-red")
    gap_msg = "Routine evidence reaches the authorization claim level." if gap_levels <= 0 else (
        "Routine evidence is one level below the authorization target." if gap_levels == 1
        else f"Routine evidence is {gap_levels} levels below the authorization target.")
    audit_items = "".join(f"<li>{x}</li>" for x in pkg["minimum_audit_dataset"])
    mfr_items = "".join(f"<li>{x}</li>" for x in pkg["manufacturer_design_requirements"])
    proc_items = "".join(f"<li>{x}</li>" for x in pkg["procurement_questions"])
    meta_parts = [today]
    if applicant and applicant not in ("unclear", ""):
        meta_parts.append(applicant)
    if submission and submission not in ("unclear", ""):
        meta_parts.append(f"Submission: {submission}")
    meta_parts.append("Grounded in 1,404 public FDA authorization records")

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>claimbounded — {device_name}</title>
<style>{_CSS_STANDALONE}</style>
</head>
<body>
<div class="cbr-wrap"><div class="cbr-container">
<header class="cbr-header">
  <div class="cbr-brand">claimbounded · Postmarket Evidence Framework · schema v3_auditability</div>
  <h1>{device_name}</h1>
  <div class="cbr-meta">{" · ".join(meta_parts)}</div>
</header>
<div class="cbr-stats">
  <div class="cbr-stat"><div class="cbr-val">{n_support}</div><div class="cbr-lbl">Supportable monitoring claims from routine data</div></div>
  <div class="cbr-stat"><div class="cbr-val">{max(gap_levels,0)}</div><div class="cbr-lbl">Levels below authorization endpoint</div></div>
  <div class="cbr-stat"><div class="cbr-val">{n_prec}</div><div class="cbr-lbl">Comparable public FDA precedents</div></div>
</div>
<div class="{callout_cls}">
  <div class="cbr-clabel">Preliminary routine-evidence claim ceiling</div>
  <div class="cbr-cvalue">{ceiling_label}</div>
  <div class="cbr-csub">{gap_msg} Re-measurable from routine data: <strong>{can_audit}</strong>.</div>
</div>
<div class="cbr-section">
  <h2>Claim Hierarchy</h2>
  <div class="cbr-ladder">{"".join(ladder)}</div>
  <p style="font-size:12px;color:#64748b;margin-top:12px">▲ Claim ceiling &nbsp;·&nbsp; ✓ Also supportable &nbsp;·&nbsp; ⚠ Authorization target above ceiling (gap) &nbsp;·&nbsp; – Not supported</p>
</div>
<div class="cbr-section">
  <h2>Claim-Support Matrix</h2>
  <div class="cbr-tw"><table>
    <thead><tr><th>Claim level</th><th>Supported by routine evidence?</th><th>Evidence required</th></tr></thead>
    <tbody>{"".join(matrix_rows)}</tbody>
  </table></div>
</div>
<div class="cbr-section">
  <h2>Gap from Authorization Endpoint</h2>
  <div class="cbr-2col">
    <div class="cbr-info"><div class="cbr-ilabel">Authorization endpoint type</div>
      <div class="cbr-ivalue">{rem.get("authorization_endpoint_type","").replace("_"," ")}</div>
      <div class="cbr-isub">Equivalent claim level: {rem.get("authorization_claim_level","").replace("_"," ")}</div></div>
    <div class="cbr-info am"><div class="cbr-ilabel">Re-measurable from routine data?</div>
      <div class="cbr-ivalue">{can_audit.upper()}</div>
      <div class="cbr-isub">{rem.get("claim_gap","")}</div></div>
  </div>
  <div class="cbr-info am"><div class="cbr-ilabel">Additional evidence work required</div>
    <div class="cbr-ivalue" style="font-size:14px;font-weight:400">{rem.get("extra_evidence_needed","")}</div></div>
</div>
<div class="cbr-section">
  <h2>Dashboard Claim Limits</h2>
  <div class="cbr-info"><div class="cbr-ilabel">Responsible monitoring claim (supportable now)</div>
    <div class="cbr-ivalue" style="font-size:14px;font-weight:400">{dash.get("responsible_dashboard_claim","")}</div></div>
  <div class="cbr-info am"><div class="cbr-ilabel">Not supportable without additional evidence</div>
    <div class="cbr-ivalue" style="font-size:14px;font-weight:400">{dash.get("not_supported_without_extra_evidence","")}</div></div>
  <div class="cbr-info"><div class="cbr-ilabel">To make the stronger claim</div>
    <div class="cbr-ivalue" style="font-size:14px;font-weight:400">{dash.get("to_make_the_stronger_claim","")}</div></div>
</div>
<div class="cbr-2col">
  <div class="cbr-section" style="margin-bottom:0"><h2>Minimum Audit Dataset</h2><ul class="cbr-list">{audit_items}</ul></div>
  <div class="cbr-section" style="margin-bottom:0"><h2>Manufacturer Design Requirements</h2><ul class="cbr-list">{mfr_items}</ul></div>
</div>
<div class="cbr-section" style="margin-top:20px"><h2>Procurement Questions</h2><ol class="cbr-olist">{proc_items}</ol></div>
<div class="cbr-section">
  <h2>Comparable Public FDA Authorization Precedents</h2>
  <div class="cbr-mnote">
    Hybrid scoring: <strong>35% regulatory identity</strong> · <strong>30% evidence structure</strong> ·
    <strong>20% text (BM25)</strong> · <strong>15% evidence-gap matching</strong>.
    All submission numbers publicly accessible at accessdata.fda.gov. Click any row for the full device profile.
  </div>
  <div class="cbr-tw"><table>
    <thead><tr><th>#</th><th>Submission</th><th>Device / Applicant</th><th>Intended use (excerpt)</th>
      <th>Endpoint type</th><th>Claim ceiling</th><th>Score</th><th>Why matched</th></tr></thead>
    <tbody>{"".join(prec_rows)}</tbody>
  </table></div>
</div>
<div class="cbr-footer">
  <p><strong>Grounding note:</strong> This preliminary classification was generated by applying the claimbounded
  study codebook (schema v3_auditability) to user-provided device inputs. The codebook was derived from
  structured extraction of 1,404 public FDA 510(k) and De Novo authorization summaries for AI-enabled medical
  devices. All precedents cited are publicly available records. This report does not constitute a regulatory
  determination, does not predict FDA decisions, and does not assess device safety or effectiveness.</p>
</div>
</div></div></body></html>"""


def _html_to_iframe(html_str: str, height: int = 820) -> str:
    encoded = base64.b64encode(html_str.encode("utf-8")).decode("utf-8")
    return (
        f'<iframe src="data:text/html;base64,{encoded}" '
        f'style="width:100%;height:{height}px;border:none;border-radius:12px;'
        f'box-shadow:0 2px 12px rgba(3,105,161,0.12);" frameborder="0"></iframe>'
    )


def _generate_docx_report(pkg: dict[str, Any]) -> str | None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    today = str(_date.today())
    cp = pkg["claim_profile"]
    device = pkg["device"]
    rem = cp["authorization_remeasurement"]
    ceiling = cp["routine_evidence_claim_ceiling"]
    ceiling_label = _CLAIM_LABELS.get(ceiling, ceiling)
    auth_claim_key = _ENDPOINT_TO_CLAIM.get(device.get("authorization_endpoint_type", ""))
    ceiling_rank = _CLAIM_RANK.get(ceiling, -1)
    gap_levels = rem.get("claim_gap_levels", 0)
    can_audit = rem.get("can_audit_authorization_endpoint_with_routine_data", "")
    device_name = device.get("device_name") or "Unnamed device"
    applicant = device.get("applicant", "")
    submission = device.get("submission_number", "")
    dash = pkg["dashboard_claim_limits"]

    def _bg(cell, color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color.lstrip("#"))
        tcPr.append(shd)

    def _run(para, text: str, bold=False, color: str | None = None,
             size: float | None = None, italic=False):
        r = para.add_run(text)
        r.bold = bold
        r.italic = italic
        if color:
            rv, gv, bv = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
            r.font.color.rgb = RGBColor(rv, gv, bv)
        if size:
            r.font.size = Pt(size)
        return r

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)

    h0 = doc.add_heading(level=0)
    _run(h0, "claimbounded", color="#0369a1", size=24, bold=True)
    h1 = doc.add_heading(level=1)
    h1.add_run(device_name)

    p = doc.add_paragraph()
    _run(p, f"Generated: {today}", color="#64748b")
    if applicant and applicant not in ("unclear", ""):
        _run(p, f"  ·  {applicant}", color="#64748b")
    if submission and submission not in ("unclear", ""):
        _run(p, f"  ·  Submission: {submission}", color="#64748b")
    p2 = doc.add_paragraph()
    _run(p2, "Grounded in 1,404 public FDA 510(k) and De Novo authorization records", italic=True, color="#64748b", size=10)
    doc.add_paragraph()

    # Claim ceiling callout
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    bg = "#dcfce7" if gap_levels <= 0 else ("#fef9c3" if gap_levels == 1 else "#fee2e2")
    fg = "#166534" if gap_levels <= 0 else ("#854d0e" if gap_levels == 1 else "#991b1b")
    _bg(cell, bg)
    cp1 = cell.paragraphs[0]
    cp1.paragraph_format.space_before = Pt(6)
    cp1.paragraph_format.space_after = Pt(2)
    _run(cp1, "PRELIMINARY ROUTINE-EVIDENCE CLAIM CEILING\n", bold=True, color="#475569", size=9)
    _run(cp1, ceiling_label, bold=True, color=fg, size=16)
    cp2 = cell.add_paragraph()
    cp2.paragraph_format.space_after = Pt(6)
    gap_msg = "Routine evidence reaches authorization claim level" if gap_levels <= 0 else f"Gap: {rem.get('claim_gap','')}"
    _run(cp2, f"{gap_msg}  ·  Re-measurable: {can_audit}", color="#475569", size=10)
    doc.add_paragraph()

    # Claim hierarchy
    doc.add_heading("Claim Hierarchy", 2)
    lt = doc.add_table(rows=len(_CLAIM_HIERARCHY), cols=1)
    lt.style = "Table Grid"
    for ri, claim in enumerate(reversed(_CLAIM_HIERARCHY)):
        rank = _CLAIM_RANK[claim]
        label = _CLAIM_LABELS[claim]
        is_ceiling = claim == ceiling
        is_gap = claim == auth_claim_key and auth_claim_key != ceiling and rank > ceiling_rank
        is_ok = rank < ceiling_rank
        cell = lt.rows[ri].cells[0]
        rp = cell.paragraphs[0]
        rp.paragraph_format.space_before = Pt(4)
        rp.paragraph_format.space_after = Pt(4)
        if is_ceiling:
            _bg(cell, "#16a34a")
            _run(rp, f"▲  {label}  —  CLAIM CEILING", bold=True, color="#ffffff", size=11)
        elif is_gap:
            _bg(cell, "#fef9c3")
            _run(rp, f"⚠  {label}  —  Authorization target (gap exists)", color="#854d0e", size=10)
        elif is_ok:
            _bg(cell, "#f0fdf4")
            _run(rp, f"✓  {label}", color="#166534", size=10)
        else:
            _bg(cell, "#f8fafc")
            _run(rp, f"–  {label}", color="#94a3b8", size=10)
    doc.add_paragraph()

    # Claim support matrix
    doc.add_heading("Claim-Support Matrix", 2)
    matrix = pkg["claim_support_matrix"]
    mt = doc.add_table(rows=1 + len(matrix), cols=3)
    mt.style = "Table Grid"
    for ci, h_text in enumerate(["Claim level", "Supported?", "Evidence required"]):
        _bg(mt.rows[0].cells[ci], "#e0f2fe")
        _run(mt.rows[0].cells[ci].paragraphs[0], h_text, bold=True, color="#0369a1", size=9)
    for ri, row in enumerate(matrix, 1):
        v = row.get("supported_by_routine_evidence", "")
        vc = "#16a34a" if v == "Yes" else ("#d97706" if "link" in v.lower() else "#94a3b8")
        mt.rows[ri].cells[0].paragraphs[0].add_run(row["claim"]).font.size = Pt(10)
        _run(mt.rows[ri].cells[1].paragraphs[0], v, bold=(v == "Yes"), color=vc, size=10)
        mt.rows[ri].cells[2].paragraphs[0].add_run(row["evidence_needed"]).font.size = Pt(9)
    doc.add_paragraph()

    # Gap analysis
    doc.add_heading("Gap from Authorization Endpoint", 2)
    for label, val, is_audit in [
        ("Authorization endpoint type", rem.get("authorization_endpoint_type","").replace("_"," "), False),
        ("Gap", rem.get("claim_gap",""), False),
        ("Re-measurable from routine data", can_audit.upper(), True),
        ("Additional evidence needed", rem.get("extra_evidence_needed",""), False),
    ]:
        rp = doc.add_paragraph()
        _run(rp, f"{label}: ", bold=True)
        color = ("#16a34a" if can_audit == "yes" else "#dc2626") if is_audit else None
        _run(rp, val, color=color, bold=is_audit)
    doc.add_paragraph()

    # Dashboard limits
    doc.add_heading("Dashboard Claim Limits", 2)
    for label, key in [
        ("Responsible monitoring claim", "responsible_dashboard_claim"),
        ("Not supportable without extra evidence", "not_supported_without_extra_evidence"),
        ("To make the stronger claim", "to_make_the_stronger_claim"),
    ]:
        rp = doc.add_paragraph()
        _run(rp, f"{label}: ", bold=True)
        _run(rp, dash.get(key, ""))
    doc.add_paragraph()

    doc.add_heading("Minimum Audit Dataset", 2)
    for item in pkg["minimum_audit_dataset"]:
        doc.add_paragraph(item, style="List Bullet").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    doc.add_heading("Manufacturer Design Requirements", 2)
    for item in pkg["manufacturer_design_requirements"]:
        doc.add_paragraph(item, style="List Bullet").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    doc.add_heading("Procurement Questions", 2)
    for item in pkg["procurement_questions"]:
        doc.add_paragraph(item, style="List Number").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # Precedents
    doc.add_heading("Comparable Public FDA Authorization Precedents", 2)
    note_p = doc.add_paragraph()
    _run(note_p, "Hybrid scoring: 35% regulatory identity · 30% evidence structure · 20% text (BM25) · 15% evidence-gap. All records publicly available at accessdata.fda.gov.", italic=True, color="#64748b", size=9)
    precs = pkg["precedents"]
    if precs:
        cols = ["#", "Submission", "Device / Applicant", "Endpoint type", "Claim ceiling", "Score", "Why matched"]
        pt = doc.add_table(rows=1 + len(precs), cols=len(cols))
        pt.style = "Table Grid"
        for ci, h_text in enumerate(cols):
            _bg(pt.rows[0].cells[ci], "#e0f2fe")
            _run(pt.rows[0].cells[ci].paragraphs[0], h_text, bold=True, color="#0369a1", size=9)
        for ri, p_item in enumerate(precs, 1):
            vals = [str(ri), p_item.get("submission_number",""),
                    f"{p_item.get('device_name','')}\n{p_item.get('applicant','')}",
                    p_item.get("authorization_endpoint_type","").replace("_"," "),
                    p_item.get("strongest_auditable_postmarket_claim","").replace("_"," "),
                    str(round(float(p_item.get("score",0)),3)), _humanize_match(p_item.get("match",""))]
            for ci, val in enumerate(vals):
                pt.rows[ri].cells[ci].paragraphs[0].add_run(val).font.size = Pt(9)
    doc.add_paragraph()

    dp = doc.add_paragraph()
    _run(dp, "Grounding note: ", bold=True, color="#475569")
    _run(dp, "This preliminary classification was generated from user-provided inputs and public FDA authorization records (schema v3_auditability). It does not constitute a regulatory determination, does not predict FDA decisions, and does not assess device safety or effectiveness.", italic=True, color="#64748b", size=9)

    tmp = tempfile.NamedTemporaryFile(suffix="_claimbounded.docx", delete=False)
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


def _format_lookup(profile: Any) -> str:
    if profile is None:
        return "_No record found. Check the submission number and try again._"
    return _format_device_panel(profile)


# ---------------------------------------------------------------------------
# Callbacks
# ---------------------------------------------------------------------------

def run_report(
    device_name, applicant, submission_number,
    submission_pathway, clinical_domain, device_function, input_data_type,
    authorization_endpoint_type, authorization_ground_truth_modality,
    routine_postmarket_evidence_stream, endpoint_linked_to_ai_output,
    endpoint_routinely_recorded, human_correction_available,
    human_overread_or_confirmation_required, mode, k,
):
    from .profiles import profile_device
    from .reports import generate_monitoring_package

    record = {
        "device_name": device_name or _UNCLEAR,
        "applicant": applicant or _UNCLEAR,
        "submission_number": submission_number or _UNCLEAR,
        "submission_pathway": submission_pathway,
        "clinical_domain": clinical_domain,
        "device_function": device_function,
        "input_data_type": input_data_type,
        "authorization_endpoint_type": authorization_endpoint_type,
        "authorization_ground_truth_modality": authorization_ground_truth_modality,
        "routine_postmarket_evidence_stream": routine_postmarket_evidence_stream,
        "endpoint_linked_to_ai_output": endpoint_linked_to_ai_output,
        "endpoint_routinely_recorded": endpoint_routinely_recorded,
        "human_correction_available": human_correction_available,
        "human_overread_or_confirmation_required": human_overread_or_confirmation_required,
    }
    profile = profile_device(record)
    pkg = generate_monitoring_package(profile, mode=mode, k=int(k))

    html_report = _generate_html_report(pkg)
    html_iframe = _html_to_iframe(html_report)

    tmp_html = tempfile.NamedTemporaryFile(mode="w", suffix="_claimbounded.html", delete=False, encoding="utf-8")
    tmp_html.write(html_report)
    tmp_html.close()

    docx_path = _generate_docx_report(pkg)

    return html_iframe, tmp_html.name, docx_path


def run_search(query, k):
    from .profiles import search_corpus

    empty = "", None, None
    if not query or not query.strip():
        return empty
    hits = search_corpus(query.strip())[: int(k)]
    if not hits:
        no_result = _html_to_iframe(
            f'<html><body style="font-family:system-ui;padding:40px;color:#0f172a">'
            f'<h2>No results for &ldquo;{query.strip()}&rdquo;</h2>'
            f'<p>Try a broader term: <code>chest</code>, <code>lung</code>, <code>stroke</code>…</p></body></html>'
        )
        return no_result, None, None

    title = f'FDA Corpus Search: "{query.strip()}" — {len(hits)} device(s)'
    html_str = _build_profile_html(hits, title=title)
    return _html_to_iframe(html_str), _save_profile_html(html_str), _generate_profile_docx(hits, title=title)


def run_lookup(submission_number):
    from .profiles import find_in_corpus

    empty = "", None, None
    if not submission_number or not submission_number.strip():
        return empty
    profile = find_in_corpus(submission_number.strip())
    if profile is None:
        not_found = _html_to_iframe(
            f'<html><body style="font-family:system-ui;padding:40px;color:#0f172a">'
            f'<h2>No record found for {submission_number.strip().upper()}</h2>'
            f'<p>Check the submission number and try again. '
            f'Use the <strong>Corpus Search</strong> tab to find valid submission numbers.</p></body></html>'
        )
        return not_found, None, None
    title = f"Device Profile: {profile.name} — {submission_number.strip().upper()}"
    html_str = _build_profile_html([profile], title=title)
    return _html_to_iframe(html_str), _save_profile_html(html_str), _generate_profile_docx([profile], title=title)


def fill_example():
    return [_EXAMPLE[f] for f in _EXAMPLE_FIELD_ORDER]


def clear_all_fields():
    return [_CLEAR[f] for f in _EXAMPLE_FIELD_ORDER]


def show_user_tip(role: str) -> str:
    tips = {
        "📋 Regulator": (
            "**You are reviewing a postmarket monitoring plan or evaluating a device's evidentiary claims.**\n\n"
            "Use the **Claim ceiling** in the report to ask: does the manufacturer's marketed monitoring claim match "
            "what their routine data can actually support? If they claim 'clinical accuracy is maintained' but the "
            "ceiling is *workflow performance*, that is a gap worth examining.\n\n"
            "The **Comparable FDA Precedents** table in the report gives you real 510(k)/De Novo submission numbers — "
            "cross-reference them on accessdata.fda.gov to see how similar devices were handled."
        ),
        "🔬 Device Manufacturer": (
            "**You are designing or improving a device's postmarket evidence architecture.**\n\n"
            "The **Manufacturer Design Requirements** section in the report is your roadmap: it tells you exactly "
            "which logging, export, and identifier features would raise your claim ceiling.\n\n"
            "The **Gap from Authorization Endpoint** section shows what additional evidence work is needed to "
            "re-measure what you demonstrated to the FDA at authorization. Use the precedent table to benchmark "
            "against comparable authorized devices."
        ),
        "🏥 Health System": (
            "**You are evaluating, deploying, or auditing an AI device in your institution.**\n\n"
            "Before deployment: use the **Procurement Questions** in the report as a vendor questionnaire. "
            "Ask vendors whether they can provide the data elements listed in the **Minimum Audit Dataset**.\n\n"
            "After deployment: the claim ceiling tells you the strongest monitoring claim your routine data "
            "supports *without* additional work. If a vendor's dashboard shows a stronger claim, treat it as "
            "unsupported until the evidence gap is closed."
        ),
    }
    return tips.get(role, "")


def on_search_select(evt, search_list):
    from .profiles import find_in_corpus
    from .schema import DeviceEvidenceProfile
    if not search_list:
        return ""
    try:
        row_idx = evt.index[0]
        if row_idx >= len(search_list):
            return ""
        d = search_list[row_idx]
        profile = find_in_corpus(d.get("submission_number", ""))
        if profile is None:
            profile = DeviceEvidenceProfile(fields=dict(d))
        return _format_device_panel(profile)
    except Exception:
        return ""


def _search_example():
    return "stroke", 20


def _search_clear():
    return "", 20


def _lookup_example():
    return "K192383"


def _lookup_clear():
    return ""


# ---------------------------------------------------------------------------
# App launcher
# ---------------------------------------------------------------------------

def launch(share: bool = False, server_port: int = 7860, server_name: str = "127.0.0.1") -> None:
    """Launch the claimbounded Gradio UI locally."""
    try:
        import gradio as gr
    except ImportError:
        raise ImportError("Gradio is required. Install with:\n    pip install claimbounded[ui]")

    DISCLAIMER = (
        "> **Disclaimer:** Preliminary classification under the study codebook from "
        "user-provided inputs and public FDA records. Not a regulatory determination. "
        "Does not assess safety, effectiveness, or FDA acceptability."
    )

    # Audience cards: solid colored backgrounds, always readable regardless of theme
    ONBOARD_HTML = """
<div style="display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin:4px 0 20px">
  <div style="background:#1565c0;border-radius:12px;padding:20px">
    <div style="font-size:22px;margin-bottom:8px">📋</div>
    <div style="font-size:15px;font-weight:700;margin-bottom:8px;color:white">Regulators</div>
    <div style="font-size:13px;color:rgba(255,255,255,0.9);line-height:1.5">
      Check whether a manufacturer's marketed monitoring claim is supportable from their
      routine data. Cross-reference real FDA submission numbers in the precedent table.
    </div>
  </div>
  <div style="background:#0f766e;border-radius:12px;padding:20px">
    <div style="font-size:22px;margin-bottom:8px">🔬</div>
    <div style="font-size:15px;font-weight:700;margin-bottom:8px;color:white">Device Manufacturers</div>
    <div style="font-size:13px;color:rgba(255,255,255,0.9);line-height:1.5">
      Get a roadmap for which logging and export features would raise your claim ceiling.
      See how comparable authorized devices were handled at FDA.
    </div>
  </div>
  <div style="background:#6b21a8;border-radius:12px;padding:20px">
    <div style="font-size:22px;margin-bottom:8px">🏥</div>
    <div style="font-size:15px;font-weight:700;margin-bottom:8px;color:white">Health Systems</div>
    <div style="font-size:13px;color:rgba(255,255,255,0.9);line-height:1.5">
      Use procurement questions as a vendor checklist. Know the strongest claim your
      routine data supports before signing a contract.
    </div>
  </div>
</div>"""

    with gr.Blocks(title="claimbounded") as demo:

        gr.Markdown("# claimbounded")
        gr.Markdown(
            "### Claim-Bounded Monitoring of AI-Enabled Medical Devices\n\n"
            "Determine what performance claim your routine deployment data can actually support — "
            "grounded in **1,404 public FDA authorization records** (510(k) and De Novo summaries)."
        )
        gr.HTML(ONBOARD_HTML)

        with gr.Tabs():

            # ══════════════════════════════════════════════════════════════
            # TAB 1 — Profile & Report
            # ══════════════════════════════════════════════════════════════
            with gr.Tab("① Profile & Report"):

                # Role selector
                user_role = gr.Radio(
                    ["📋 Regulator", "🔬 Device Manufacturer", "🏥 Health System"],
                    label="Select your role for tailored guidance (optional)",
                    value=None,
                )
                tip_box = gr.Markdown()
                user_role.change(fn=show_user_tip, inputs=[user_role], outputs=[tip_box])

                gr.Markdown(
                    "Fill in what you know about the device. Fields are pre-filled with a worked example "
                    "(LVO triage). Hit **Clear All** to start fresh, or **Auto-complete Example** to reset."
                )

                # Buttons BEFORE the form
                with gr.Row():
                    example_btn = gr.Button("✦ Auto-complete Example", variant="secondary")
                    clear_btn = gr.Button("✕ Clear All Fields", variant="secondary")

                # ── Form ─────────────────────────────────────────────────
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("##### Device identity")
                        inp_device_name = gr.Textbox(label="Device name", value=_EXAMPLE["device_name"], placeholder="e.g. Acme LVO Triage")
                        inp_applicant = gr.Textbox(label="Manufacturer / applicant", value=_EXAMPLE["applicant"], placeholder="e.g. Acme Medical Inc.")
                        inp_submission = gr.Textbox(label="Submission number (if known)", value="", placeholder="e.g. K192383")
                        inp_pathway = gr.Dropdown(SUBMISSION_PATHWAYS, value=_EXAMPLE["submission_pathway"], label="Submission pathway",
                            info="510(k) = substantial equivalence to a predicate device. De Novo = novel low-to-moderate risk device without a predicate.")
                        inp_domain = gr.Dropdown(CLINICAL_DOMAINS, value=_EXAMPLE["clinical_domain"], label="Clinical domain",
                            info="Primary medical specialty of the device's intended use.")
                        gr.Markdown("##### Technical profile")
                        inp_function = gr.Dropdown(DEVICE_FUNCTIONS, value=_EXAMPLE["device_function"], label="Device function",
                            info="triage_notification = flags urgent cases · diagnostic_classification = classifies findings · segmentation_quantification = delineates and measures structures · risk_prediction_prognosis = estimates future events · physiologic_monitoring = monitors signals.")
                        inp_input_type = gr.Dropdown(INPUT_DATA_TYPES, value=_EXAMPLE["input_data_type"], label="Input data type",
                            info="Type of data the device processes (radiology = CT/MRI/X-ray, pathology = slide images, physiologic = ECG/EEG).")
                        inp_endpoint_type = gr.Dropdown(ENDPOINT_TYPES, value=_EXAMPLE["authorization_endpoint_type"], label="Authorization endpoint type",
                            info="diagnostic_accuracy = sensitivity/specificity/AUC · triage_sensitivity_specificity = binary detection · quantitative_measurement_agreement = Bland-Altman/ICC · segmentation_geometric_accuracy = Dice/Hausdorff · workflow_or_time_to_notification = delivery time.")

                    with gr.Column(scale=1):
                        gr.Markdown("##### Evidence profile")
                        inp_ground_truth = gr.Dropdown(GROUND_TRUTH_MODALITIES, value=_EXAMPLE["authorization_ground_truth_modality"],
                            label="Authorization ground truth modality",
                            info="expert_reader_panel = adjudication by multiple experts · clinical_diagnosis = diagnosis from medical record · laboratory_reference_method = gold-standard lab assay · pathology_or_histology = tissue examination · longitudinal_clinical_outcome = confirmed outcome at follow-up.")
                        inp_stream = gr.Dropdown(EVIDENCE_STREAMS, value=_EXAMPLE["routine_postmarket_evidence_stream"],
                            label="Routine postmarket evidence stream",
                            info="workflow_logs = timestamps/delivery/acknowledgement · human_corrections_or_edits = recorded edits to AI outputs · clinician_acceptance_or_override = accept/reject/override decisions · output_logs_only = AI scores with no downstream linkage · none_described = no evidence stream documented.")
                        gr.Markdown("##### Deployment evidence — key ceiling drivers")
                        inp_linked = gr.Radio(YES_NO_UNCLEAR, value=_EXAMPLE["endpoint_linked_to_ai_output"],
                            label="AI output linked case-level to downstream reference?",
                            info="Is each AI output individually matched to a downstream reference result or clinical report?")
                        inp_recorded = gr.Radio(YES_NO_UNCLEAR, value=_EXAMPLE["endpoint_routinely_recorded"],
                            label="Endpoint routinely recorded in clinical workflow?",
                            info="Is the authorization endpoint outcome documented in routine clinical records without extra effort?")
                        inp_correction = gr.Radio(YES_NO_UNCLEAR, value=_EXAMPLE["human_correction_available"],
                            label="Clinician accept / edit / override captured on AI output?",
                            info="Does the system log whether a clinician accepted, modified, or rejected each AI recommendation?")
                        inp_overread = gr.Radio(YES_NO_UNCLEAR, value=_EXAMPLE["human_overread_or_confirmation_required"],
                            label="Human overread or confirmation required before clinical use?",
                            info="Does clinical policy require a human to review the AI output before it influences care decisions?")
                        gr.Markdown("##### Precedent retrieval")
                        inp_mode = gr.Radio(RETRIEVAL_MODES, value="hybrid", label="Retrieval mode",
                            info="hybrid = best overall (recommended) · like_for_like = same regulatory identity · adjacent = same evidence problem · claim_gap = same authorization-ceiling divergence.")
                        inp_k = gr.Slider(1, 20, value=8, step=1, label="Number of precedents (k)")

                all_inputs = [
                    inp_device_name, inp_applicant, inp_submission,
                    inp_pathway, inp_domain, inp_function, inp_input_type, inp_endpoint_type,
                    inp_ground_truth, inp_stream,
                    inp_linked, inp_recorded, inp_correction, inp_overread,
                ]
                example_btn.click(fn=fill_example, inputs=[], outputs=all_inputs)
                clear_btn.click(fn=clear_all_fields, inputs=[], outputs=all_inputs)

                # Generate button after the form
                run_btn = gr.Button("Generate Report", variant="primary", size="lg")

                # ── Results ───────────────────────────────────────────────
                gr.Markdown("---")
                gr.Markdown(
                    "### Monitoring Report\n\n"
                    "The full report renders below — including the claim hierarchy, gap analysis, "
                    "dashboard limits, audit dataset, and comparable FDA precedents. "
                    "Download as HTML or Word for offline use and sharing."
                )

                out_html_view = gr.HTML()

                with gr.Row():
                    out_html_dl = gr.DownloadButton(label="⬇ Download HTML Report", variant="secondary")
                    out_docx_dl = gr.DownloadButton(label="⬇ Download Word Document (.docx)", variant="primary")

                gr.Markdown(DISCLAIMER)

                run_btn.click(
                    fn=run_report,
                    inputs=[
                        inp_device_name, inp_applicant, inp_submission,
                        inp_pathway, inp_domain, inp_function, inp_input_type, inp_endpoint_type,
                        inp_ground_truth, inp_stream,
                        inp_linked, inp_recorded, inp_correction, inp_overread,
                        inp_mode, inp_k,
                    ],
                    outputs=[out_html_view, out_html_dl, out_docx_dl],
                )

            # ══════════════════════════════════════════════════════════════
            # TAB 2 — Corpus Search
            # ══════════════════════════════════════════════════════════════
            with gr.Tab("② Corpus Search"):
                gr.Markdown(
                    "Search the **1,404-device FDA corpus** by device name, manufacturer, or intended use. "
                    "Results render as a full stakeholder report — download as HTML or Word.\n\n"
                    "**Try:** `diabetic retinopathy` · `lung nodule` · `sepsis` · "
                    "`mammography` · `arrhythmia` · `polyp` · `stroke` · `fracture`"
                )
                with gr.Row():
                    search_example_btn = gr.Button("✦ Example: stroke", variant="secondary")
                    clear_search_btn = gr.Button("✕ Clear", variant="secondary")
                with gr.Row():
                    inp_search_query = gr.Textbox(
                        label="Search query", placeholder="e.g. large vessel occlusion", scale=4
                    )
                    inp_search_k = gr.Slider(1, 100, value=20, step=1, label="Max results", scale=1)
                search_btn = gr.Button("Search", variant="primary")

                out_search_view = gr.HTML()
                with gr.Row():
                    out_search_html = gr.DownloadButton(label="⬇ Download HTML Report", variant="secondary")
                    out_search_docx = gr.DownloadButton(label="⬇ Download Word Document (.docx)", variant="primary")
                gr.Markdown(DISCLAIMER)

                search_btn.click(
                    fn=run_search,
                    inputs=[inp_search_query, inp_search_k],
                    outputs=[out_search_view, out_search_html, out_search_docx],
                )
                search_example_btn.click(fn=_search_example, inputs=[], outputs=[inp_search_query, inp_search_k])
                clear_search_btn.click(fn=_search_clear, inputs=[], outputs=[inp_search_query, inp_search_k])

            # ══════════════════════════════════════════════════════════════
            # TAB 3 — Submission Lookup
            # ══════════════════════════════════════════════════════════════
            with gr.Tab("③ Submission Lookup"):
                gr.Markdown(
                    "Look up a specific FDA submission by number (e.g. `K192383`). "
                    "Renders a full stakeholder profile — download as HTML or Word.\n\n"
                    "**Tip:** Find submission numbers via Corpus Search, or from the precedents "
                    "table in the downloaded HTML/Word report from Tab 1."
                )
                with gr.Row():
                    lookup_example_btn = gr.Button("✦ Example: K192383", variant="secondary")
                    clear_lookup_btn = gr.Button("✕ Clear", variant="secondary")
                with gr.Row():
                    inp_lookup = gr.Textbox(label="Submission number", placeholder="K192383", scale=3)
                    lookup_btn = gr.Button("Lookup", variant="primary", scale=1)

                out_lookup_view = gr.HTML()
                with gr.Row():
                    out_lookup_html = gr.DownloadButton(label="⬇ Download HTML Profile", variant="secondary")
                    out_lookup_docx = gr.DownloadButton(label="⬇ Download Word Document (.docx)", variant="primary")
                gr.Markdown(DISCLAIMER)

                lookup_btn.click(
                    fn=run_lookup,
                    inputs=[inp_lookup],
                    outputs=[out_lookup_view, out_lookup_html, out_lookup_docx],
                )
                lookup_example_btn.click(fn=_lookup_example, inputs=[], outputs=[inp_lookup])
                clear_lookup_btn.click(fn=_lookup_clear, inputs=[], outputs=[inp_lookup])

    print()
    print("=" * 56)
    print("  claimbounded UI launching...")
    print(f"  → http://localhost:{server_port}")
    print("  All processing runs locally. No data leaves your machine.")
    print("=" * 56)
    print()

    demo.launch(inbrowser=True, share=share, server_port=server_port,
                server_name=server_name, theme=gr.themes.Soft())


if __name__ == "__main__":  # pragma: no cover
    launch()
